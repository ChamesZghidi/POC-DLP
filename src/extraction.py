"""Extraction de contenu pour le pipeline DLP.

Les archives ne sont jamais extraites sur disque : cela évite les traversées de
chemin et limite l'impact des archives volumineuses ou malveillantes.
"""

from __future__ import annotations

import email
from email import policy
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import zipfile
from xml.etree import ElementTree as ET

import docx
from pypdf import PdfReader


MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 200
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 30 * 1024 * 1024
MAX_NESTED_ARCHIVE_DEPTH = 2
TEXT_EXTENSIONS = {".txt", ".csv", ".json", ".log", ".md"}


def _metadata(**values: object) -> dict:
    values.setdefault("est_email", False)
    return values


def extract_from_docx(file_path_or_bytes: str | bytes | BytesIO) -> dict:
    """Extrait les paragraphes et tableaux d'un fichier Word."""
    document = docx.Document(file_path_or_bytes)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return {"text": "\n".join(paragraphs), "metadata": _metadata(format="docx")}


def extract_from_pdf(file_path_or_bytes: str | bytes | BytesIO) -> dict:
    reader = PdfReader(file_path_or_bytes)
    return {
        "text": "\n".join(page.extract_text() or "" for page in reader.pages),
        "metadata": _metadata(format="pdf", pages=len(reader.pages)),
    }


def extract_from_eml(file_path_or_bytes: str | bytes | BytesIO) -> dict:
    if isinstance(file_path_or_bytes, (str, Path)):
        with open(file_path_or_bytes, "rb") as source:
            raw = source.read()
    elif hasattr(file_path_or_bytes, "read"):
        raw = file_path_or_bytes.read()
    else:
        raw = file_path_or_bytes
    msg = email.message_from_bytes(raw, policy=policy.default)
    body = ""
    for part in msg.walk() if msg.is_multipart() else [msg]:
        disposition = str(part.get("Content-Disposition", ""))
        if "attachment" in disposition.lower():
            continue
        if part.get_content_type() in {"text/plain", "text/html"}:
            payload = part.get_payload(decode=True) or b""
            value = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
            body = re.sub(r"<[^>]+>", " ", value) if part.get_content_type() == "text/html" else value
            break
    subject, sender, recipient, date = (str(msg.get(key, "")) for key in ("Subject", "From", "To", "Date"))
    text = f"Sujet: {subject}\nDe: {sender}\nÀ: {recipient}\nDate: {date}\n\n{body.strip()}"
    return {"text": unescape(text), "metadata": _metadata(format="eml", est_email=True, sujet=subject,
            expediteur=sender, destinataire=recipient, date=date)}


def extract_from_text(raw: bytes) -> dict:
    return {"text": raw.decode("utf-8", errors="replace"), "metadata": _metadata(format="text")}


def extract_from_xml(raw: bytes) -> dict:
    """Extrait uniquement le texte XML, sans résoudre d'entités externes."""
    if b"<!DOCTYPE" in raw.upper() or b"<!ENTITY" in raw.upper():
        raise ValueError("XML avec DTD/entités non autorisé")
    try:
        root = ET.fromstring(raw)
    except ET.ParseError as exc:
        raise ValueError(f"XML invalide : {exc}") from exc
    text = " ".join(fragment.strip() for fragment in root.itertext() if fragment.strip())
    return {"text": text, "metadata": _metadata(format="xml", xml_root=root.tag)}


def _extract_bytes(raw: bytes, extension: str, *, depth: int = 0) -> dict:
    if len(raw) > MAX_FILE_BYTES and extension != ".zip":
        raise ValueError("Fichier trop volumineux pour l'analyse DLP")
    extension = extension.lower()
    if extension == ".docx":
        return extract_from_docx(BytesIO(raw))
    if extension == ".pdf":
        return extract_from_pdf(BytesIO(raw))
    if extension == ".eml":
        return extract_from_eml(raw)
    if extension == ".xml":
        return extract_from_xml(raw)
    if extension in TEXT_EXTENSIONS:
        return extract_from_text(raw)
    if extension == ".zip":
        return extract_from_zip(raw, depth=depth)
    raise ValueError(f"Format non supporté : {extension or '(sans extension)'}")


def extract_from_zip(file_path_or_bytes: str | bytes, *, depth: int = 0) -> dict:
    """Analyse de façon bornée les fichiers supportés présents dans une archive ZIP."""
    if depth >= MAX_NESTED_ARCHIVE_DEPTH:
        raise ValueError("Profondeur maximale d'archives imbriquées atteinte")
    raw = Path(file_path_or_bytes).read_bytes() if isinstance(file_path_or_bytes, (str, Path)) else file_path_or_bytes
    if len(raw) > MAX_FILE_BYTES:
        raise ValueError("Archive ZIP trop volumineuse")
    try:
        archive = zipfile.ZipFile(BytesIO(raw))
    except zipfile.BadZipFile as exc:
        raise ValueError("Archive ZIP invalide") from exc

    infos = [info for info in archive.infolist() if not info.is_dir()]
    if len(infos) > MAX_ARCHIVE_ENTRIES:
        raise ValueError("Archive ZIP : trop de fichiers")
    declared_size = sum(info.file_size for info in infos)
    if declared_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
        raise ValueError("Archive ZIP : taille décompressée trop importante")

    chunks, analysed, skipped, errors, analysis_items = [], [], [], [], []
    for info in infos:
        name = info.filename.replace("\\", "/")
        suffix = Path(name).suffix.lower()
        if suffix not in TEXT_EXTENSIONS | {".docx", ".pdf", ".eml", ".xml", ".zip"}:
            skipped.append(name)
            continue
        try:
            item = _extract_bytes(archive.read(info), suffix, depth=depth + 1)
            if item["text"].strip():
                chunks.append(f"\n--- Fichier archive : {name} ---\n{item['text']}")
                # A ZIP is a container, not one document. Keep member
                # boundaries so a finance word cannot mask a medical file.
                analysis_items.append({"name": name, "text": item["text"]})
            analysed.append(name)
        except (ValueError, zipfile.BadZipFile, OSError) as exc:
            errors.append({"file": name, "error": str(exc)})
    return {"text": "\n".join(chunks).strip(), "metadata": _metadata(format="zip", archive_entries=len(infos),
            analysed_files=analysed, skipped_files=skipped, extraction_errors=errors,
            analysis_items=analysis_items)}


def extract_text(file_path: str) -> dict:
    """Point d'entrée unique pour DOCX, PDF, EML, XML, ZIP et texte."""
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"Fichier introuvable : {path}")
    result = _extract_bytes(path.read_bytes(), path.suffix.lower())
    result["metadata"].update({"file_name": path.name, "source_extension": path.suffix.lower()})
    return result
